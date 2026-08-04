"""AI 신청서(사업계획서) 작성 도우미 - 순수 로직 모듈 (Streamlit UI 코드 없음).
공고문+신청서양식 분석 -> 항목별 초안 생성 -> 채팅으로 초안 수정 -> 저장/문서 출력까지
지원한다. matcher.py/chatbot.py와 동일하게 Gemini 호출과 Supabase 접근만 담당하고,
화면 렌더링은 pages/3_신청서작성.py에서 처리한다.
"""

import io
import json
import os
import re
from datetime import datetime, timezone

from google import genai
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import CONTENT_TYPE as OPC_CONTENT_TYPE, RELATIONSHIP_TYPE as RT
from docx.opc.packuri import PackURI
from docx.opc.part import Part
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.text.paragraph import Paragraph
from dotenv import load_dotenv
from lxml import etree
import requests

import matcher

load_dotenv(override=True)

GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
ai_client = genai.Client(api_key=GEMINI_API_KEY)

# 요건분석처럼 짧고 구조화된 추출 작업은 저렴한 lite 모델로 충분하지만,
# 실제 "글쓰기"(초안 생성/채팅 수정)는 분량과 구체성이 중요해 상위 모델을 쓴다.
ANALYSIS_MODEL = "gemini-flash-lite-latest"
WRITING_MODEL = "gemini-flash-latest"

REQUIREMENTS_PROMPT_TEMPLATE = """
당신은 대한민국 정부 지원사업 신청서 작성 컨설턴트입니다. 아래 공고문 및 신청서 양식 내용을 분석하여,
신청서 작성 시 채워야 할 항목과 심사 시 중요하게 보는 핵심 요소, 준비해야 할 서류를 지정된 JSON
형식으로 정리하세요. 신청서 양식 원문이 없거나 부실하면, 정부지원사업 신청서에 통상적으로 들어가는
항목(사업 개요, 추진 배경/필요성, 추진 계획, 기대효과, 소요예산 등)을 공고 내용에 맞게 구성하세요.

[공고문 내용]
{announcement_text}

[신청서 양식 원문 (없으면 "(없음)")]
{form_text}

[추출할 JSON 스키마]
{{
  "form_sections": [{{"section_name": "항목명", "guidance": "이 항목에 무엇을, 어떤 관점으로 써야 하는지 구체적 안내"}}],
  "evaluation_criteria": ["심사/선정 시 중요하게 보는 핵심 요소"],
  "required_documents": ["신청 시 제출이 필요한 준비서류 목록"]
}}
"""

DRAFT_PROMPT_TEMPLATE = """
당신은 대한민국 정부 지원사업 신청서 작성을 돕는 컨설턴트입니다. 아래 기업 정보와 추가 자료, 그리고
신청서 항목별 작성 가이드와 심사 핵심요소를 참고하여 각 항목의 초안을 작성하세요. 기업 정보에 없는
구체적 수치나 실적을 지어내지 말고 "[확인 필요: ...]" 형태로 표시하세요. 심사 핵심요소를 자연스럽게
반영하되 과장하지 마세요.

분량과 구체성 기준:
- 각 항목은 최소 500자 이상, 실제 심사위원이 읽을 수준으로 충실하게 작성하세요. 한두 문장으로
  요약하듯 쓰지 말고, 배경-현황-구체적 실행방안-근거 순으로 문단을 구성하세요.
- "혁신적인", "다양한", "효과적으로" 같은 추상적 수사보다, 추가 자료에 있는 구체적 수치·일정·
  대상·방법을 우선 활용하세요. 추가 자료에 없는 수치는 지어내지 말고 [확인 필요: 항목]으로 표시하되,
  그 주변 서술(맥락, 방법론, 논리)은 최대한 구체적으로 채우세요.
- 추진 계획류 항목은 월차/분기별 마일스톤처럼 시간 순서가 드러나게, 예산/기대효과류 항목은 항목별
  세부 내역이 드러나게 작성하세요.

[기업 정보]
{company_profile_json}

[추가 자료 (사업 내용, 실적 등)]
{extra_context}

[작성해야 할 항목 및 가이드]
{form_sections_json}

[심사 핵심요소]
{evaluation_criteria_json}

[출력 형식]
반드시 아래 JSON 형식으로만 답하세요. 키는 위 "작성해야 할 항목"의 section_name과 정확히 일치시키세요.
{{"<section_name>": "<초안 본문>", "...": "..."}}
"""

REFINE_CHAT_PROMPT_TEMPLATE = """
당신은 대한민국 정부 지원사업 신청서 작성을 돕는 컨설턴트입니다. 사용자가 이미 작성된 신청서 초안에
대해 수정이나 정보 추가를 요청하면, 관련된 항목만 다시 작성해서 반영하세요. 사용자가 요청하지 않은
항목은 절대 건드리지 마세요.

수정할 때도 분량을 줄이지 마세요 - 요청한 내용을 기존 문단에 자연스럽게 녹여 넣어 최소 500자 이상,
심사위원이 읽을 수준의 완결된 문단으로 다시 쓰세요. 단순히 한 문장만 덧붙이는 식으로 답하지 마세요.

[기업 정보]
{company_profile_json}

[추가 자료]
{extra_context}

[심사 핵심요소]
{evaluation_criteria_json}

[현재 신청서 초안 (항목별)]
{sections_json}

[이전 대화]
{history}

[사용자의 이번 요청]
{user_message}

[출력 형식]
{{
  "reply": "무엇을 어떻게 반영했는지 사용자에게 보여줄 간결한 한국어 설명",
  "updated_sections": {{"수정한 항목명": "새 본문", "...": "..."}}
}}
사용자 요청이 특정 항목과 무관한 일반 질문/답변이면 updated_sections는 빈 객체({{}})로 두고 reply에만 답하세요.
"""

TEMPLATE_MAP_PROMPT = """
아래는 워드(.docx) 신청서 양식에서 내용을 채워 넣을 수 있는 '위치 후보' 목록입니다. 각 후보는
번호(index)와 종류(kind), 그리고 무엇을 채워야 할 자리인지 알려주는 라벨(label)로 구성됩니다.
- kind가 "paragraph"인 경우: 그 문단(제목/라벨/안내문) 바로 다음에 내용이 들어갑니다.
- kind가 "table_cell"인 경우: 한국 정부 신청서에 매우 흔한 [라벨 칸 | 값 칸] 표 구조에서, label에
  적힌 라벨 칸 바로 옆(값 칸)에 내용이 들어갑니다. 표 구조인 양식은 이 표 셀 후보를 우선 활용하세요.

그리고 신청서 초안의 작성 항목 목록이 있습니다. 각 작성 항목의 내용이 들어가기에 가장 자연스러운
위치 후보를 찾아 매핑하세요. 라벨의 표현이 완전히 같지 않아도 의미상 대응되면 매핑하세요
(예: 항목명 "지원 동기" ↔ 라벨 "1. 신청 배경 및 수출 필요성"). 해당하는 후보를 찾을 수 없으면
-1로 표시하세요. 서로 다른 항목을 같은 후보에 매핑하지 마세요.

[위치 후보 목록]
{targets_json}

[작성 항목 목록]
{section_names_json}

[출력 형식]
반드시 아래 JSON 형식으로만 답하세요. 키는 위 "작성 항목 목록"의 이름과 정확히 일치시키세요.
{{"<항목명>": <후보 번호 또는 -1>, "...": ...}}
"""


def _format_chat_history(history: list, max_turns: int = 6) -> str:
    if not history:
        return "(없음)"
    recent = history[-(max_turns * 2):]
    return "\n".join(f"{'사용자' if m['role'] == 'user' else '컨설턴트'}: {m['content']}" for m in recent)


def _parse_json_response(response) -> dict:
    """가벼운 모델이 가끔 객체를 배열로 감싸거나 형식을 어기는 경우를 보정한다."""
    try:
        parsed = json.loads(response.text.strip())
    except (json.JSONDecodeError, AttributeError):
        return {}
    if isinstance(parsed, list):
        parsed = parsed[0] if parsed and isinstance(parsed[0], dict) else {}
    return parsed if isinstance(parsed, dict) else {}


def extract_application_requirements(
    announcement_text: str = "",
    form_text: str = "",
    announcement_images: list | None = None,
    form_images: list | None = None,
) -> dict:
    prompt = REQUIREMENTS_PROMPT_TEMPLATE.format(
        announcement_text=(announcement_text or "")[:6000] or "(없음)",
        form_text=(form_text or "")[:6000] or "(없음)",
    )
    images = (announcement_images or []) + (form_images or [])
    contents = [prompt] + images if images else prompt

    response = ai_client.models.generate_content(
        model=ANALYSIS_MODEL,
        contents=contents,
        config={"response_mime_type": "application/json"},
    )
    parsed = _parse_json_response(response)
    parsed.setdefault("form_sections", [])
    parsed.setdefault("evaluation_criteria", [])
    parsed.setdefault("required_documents", [])
    return parsed


def draft_application_sections(company_profile: dict, extra_context: str, requirements: dict) -> dict:
    sections = requirements.get("form_sections") or []
    if not sections:
        return {}

    prompt = DRAFT_PROMPT_TEMPLATE.format(
        company_profile_json=json.dumps(company_profile, ensure_ascii=False),
        extra_context=extra_context or "(없음)",
        form_sections_json=json.dumps(sections, ensure_ascii=False),
        evaluation_criteria_json=json.dumps(requirements.get("evaluation_criteria") or [], ensure_ascii=False),
    )
    response = ai_client.models.generate_content(
        model=WRITING_MODEL,
        contents=prompt,
        config={"response_mime_type": "application/json"},
    )
    drafted = _parse_json_response(response)

    # AI가 일부 항목을 누락해도 화면에 빈 칸으로라도 항상 표시되도록 보정
    for s in sections:
        name = s.get("section_name")
        if name and name not in drafted:
            drafted[name] = ""
    return drafted


def refine_draft_via_chat(
    sections: dict,
    requirements: dict,
    company_profile: dict,
    extra_context: str,
    chat_history: list,
    user_message: str,
) -> dict:
    prompt = REFINE_CHAT_PROMPT_TEMPLATE.format(
        company_profile_json=json.dumps(company_profile, ensure_ascii=False),
        extra_context=extra_context or "(없음)",
        evaluation_criteria_json=json.dumps(requirements.get("evaluation_criteria") or [], ensure_ascii=False),
        sections_json=json.dumps(sections, ensure_ascii=False),
        history=_format_chat_history(chat_history),
        user_message=user_message,
    )
    response = ai_client.models.generate_content(
        model=WRITING_MODEL,
        contents=prompt,
        config={"response_mime_type": "application/json"},
    )
    parsed = _parse_json_response(response)
    parsed.setdefault("reply", "요청을 처리하지 못했습니다. 다시 시도해 주세요.")
    updated = parsed.get("updated_sections")
    parsed["updated_sections"] = updated if isinstance(updated, dict) else {}
    return parsed


def save_draft(
    draft_id: str | None,
    company_name: str,
    announcement_title: str,
    announcement_detail_url: str,
    requirements: dict,
    draft_sections: dict,
    company_profile: dict | None = None,
    extra_context: str = "",
) -> str:
    """draft_id가 없으면 새로 만들고, 있으면 같은 행을 덮어쓴다. 저장된(또는 갱신된) 행의 id를 반환.
    company_profile/extra_context까지 함께 저장해두면, 나중에 초안을 불러와 채팅으로 계속
    이어서 수정할 때도 AI가 같은 맥락(회사 정보, 추가 자료)을 참고할 수 있다."""
    row = {
        "company_name": company_name,
        "announcement_title": announcement_title,
        "announcement_detail_url": announcement_detail_url,
        "requirements": requirements,
        "draft_sections": draft_sections,
        "company_profile": company_profile or {},
        "extra_context": extra_context or "",
        "updated_at": datetime.now(timezone.utc).isoformat(),
    }
    if draft_id:
        row["id"] = draft_id

    response = matcher.supabase.table("application_drafts").upsert(row).execute()
    return response.data[0]["id"] if response.data else draft_id


def load_saved_drafts() -> list:
    response = (
        matcher.supabase.table("application_drafts")
        .select("*")
        .order("updated_at", desc=True)
        .execute()
    )
    return response.data


DOCX_NAVY = RGBColor(0x1F, 0x3A, 0x63)
DOCX_LIGHT_BLUE = "EAF1FB"
DOCX_PLACEHOLDER_COLOR = RGBColor(0xC0, 0x39, 0x2B)
_PLACEHOLDER_RE = re.compile(r"(\[확인\s*필요[^\]]*\])")


def _set_cell_background(cell, hex_color: str):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def _add_left_accent_border(paragraph, color_hex: str = "1F3A63"):
    """섹션 제목 왼쪽에 컬러 바를 둬서(참고 문서 스타일) 구획을 한눈에 구분되게 한다."""
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "24")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), color_hex)
    p_bdr.append(left)
    p_pr.append(p_bdr)


def _add_body_text(paragraph, text: str, comments: list):
    """'[확인 필요: ...]' 표시는 빨간 굵은 글씨로 강조해 본문만 봐도 바로 눈에 띄게 하고,
    동시에 실제 Word 댓글(주석)로도 anchor해 검토자가 댓글 패널에서 하나씩 확인하며
    지워나갈 수 있게 한다. comments 리스트에 (댓글 본문)을 순서대로 쌓아두면,
    문서 전체를 다 만든 뒤 _attach_comments()가 실제 comments.xml 파트로 묶어 붙인다."""
    for part in _PLACEHOLDER_RE.split(text):
        if not part:
            continue
        run = paragraph.add_run(part)
        if _PLACEHOLDER_RE.match(part):
            run.bold = True
            run.font.color.rgb = DOCX_PLACEHOLDER_COLOR
            inner = re.sub(r"^확인\s*필요\s*:?\s*", "", part.strip("[]"))
            comment_id = len(comments)
            comments.append(f"⚠️ 검토 필요 — 실제 정보로 채워 넣어 주세요: {inner}")
            _anchor_comment(run, comment_id)


def _anchor_comment(run, comment_id: int):
    """run 하나를 commentRangeStart/End + commentReference로 감싸, comments.xml의
    같은 id를 가진 댓글과 연결한다 (python-docx엔 댓글 API가 없어 원시 OOXML로 직접 구성)."""
    start = OxmlElement("w:commentRangeStart")
    start.set(qn("w:id"), str(comment_id))
    end = OxmlElement("w:commentRangeEnd")
    end.set(qn("w:id"), str(comment_id))

    ref_run = OxmlElement("w:r")
    ref_rpr = OxmlElement("w:rPr")
    rstyle = OxmlElement("w:rStyle")
    rstyle.set(qn("w:val"), "CommentReference")
    ref_rpr.append(rstyle)
    ref_run.append(ref_rpr)
    ref = OxmlElement("w:commentReference")
    ref.set(qn("w:id"), str(comment_id))
    ref_run.append(ref)

    run._r.addprevious(start)
    run._r.addnext(end)
    end.addnext(ref_run)


def _attach_comments(doc: Document, comment_texts: list[str]):
    """수집된 댓글 목록을 실제 word/comments.xml 파트로 만들어 문서 패키지에 연결한다.
    python-docx는 댓글을 직접 다루는 API가 없어서, OPC(Open Packaging Convention) 레벨에서
    파트를 만들고 관계(relationship)를 맺어주면 된다 - 그러면 [Content_Types].xml과
    document.xml.rels는 python-docx가 저장 시점에 알아서 채워준다."""
    if not comment_texts:
        return

    comments_el = OxmlElement("w:comments")
    now = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")
    for i, text in enumerate(comment_texts):
        comment = OxmlElement("w:comment")
        comment.set(qn("w:id"), str(i))
        comment.set(qn("w:author"), "AI 초안 검토")
        comment.set(qn("w:date"), now)
        comment.set(qn("w:initials"), "AI")
        p = OxmlElement("w:p")
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = text
        r.append(t)
        p.append(r)
        comment.append(p)
        comments_el.append(comment)

    blob = etree.tostring(comments_el, xml_declaration=True, encoding="UTF-8", standalone=True)
    comments_part = Part(
        PackURI("/word/comments.xml"), OPC_CONTENT_TYPE.WML_COMMENTS, blob, doc.part.package
    )
    doc.part.relate_to(comments_part, RT.COMMENTS)


def _lock_table_layout(table, total_width: "Cm"):
    """표 너비를 '고정'으로 잠그고 전체 너비를 명시한다. python-docx 기본값인 자동맞춤
    (autofit)은 Word에서는 대체로 괜찮지만 Google Docs 등 다른 뷰어에서 셀 너비 지정을
    무시하고 내용 기준으로 다시 배치해버리는 경우가 있어, 뷰어에 관계없이 레이아웃이 일정
    하게 나오도록 고정폭으로 강제한다. python-docx의 Table에는 width 속성이 없어(1.2.0
    기준) tblW를 OOXML로 직접 써야 한다."""
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(total_width.twips))


def _style_or_none(doc: Document, style_name: str):
    """이름으로 스타일을 찾되, 없으면 None을 반환한다 (add_paragraph(style=None)은 '기본'
    스타일이 되어 에러 없이 동작함). 우리가 새로 만드는 문서(build_docx)는 python-docx의
    기본 템플릿을 쓰므로 'Heading 2'가 항상 있지만, 사용자가 올리거나 공고에서 내려받은
    기존 .docx 양식(fill_docx_template)은 정의된 스타일이 제각각이라 없을 수 있다."""
    try:
        return doc.styles[style_name]
    except KeyError:
        return None


def _add_page_number_footer(doc: Document):
    footer_p = doc.sections[0].footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p.add_run()
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def build_docx(announcement: dict, company_profile: dict, sections: dict) -> io.BytesIO:
    """참고 문서(그라데이션 배너 + 카드형 섹션) 스타일을 python-docx로 재현한 신청서 초안.
    - 파란 제목 배너 + 메타정보 표로 한눈에 어떤 공고/기업용 초안인지 보이게 하고
    - 섹션마다 왼쪽 컬러 바로 구획을 나누고
    - 미확정 수치는 빨간 글씨로 강조해 제출 전 검토 지점을 표시한다.
    """
    announcement = announcement or {}
    company_profile = company_profile or {}
    title = announcement.get("title") or "신청서 초안"

    doc = Document()
    for sec in doc.sections:
        sec.left_margin = Cm(2)
        sec.right_margin = Cm(2)

    normal = doc.styles["Normal"]
    normal.font.name = "맑은 고딕"
    normal.font.size = Pt(10.5)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")

    # 제목 배너
    banner = doc.add_table(rows=1, cols=1)
    banner.alignment = WD_TABLE_ALIGNMENT.CENTER
    _lock_table_layout(banner, Cm(16.7))
    banner_cell = banner.rows[0].cells[0]
    _set_cell_background(banner_cell, "1F3A63")
    title_p = banner_cell.paragraphs[0]
    title_run = title_p.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(19)
    title_run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    sub_p = banner_cell.add_paragraph()
    sub_run = sub_p.add_run("AI 신청서 작성 도우미로 생성된 초안 문서")
    sub_run.font.size = Pt(9.5)
    sub_run.font.color.rgb = RGBColor(0xCE, 0xDC, 0xEE)

    doc.add_paragraph()

    # 메타정보 표
    meta_rows = [
        ("소관기관", announcement.get("department") or "정보 없음"),
        ("신청기업", company_profile.get("company_name") or "정보 없음"),
        ("공고 원문", announcement.get("detail_url") or "-"),
        ("작성일", datetime.now().strftime("%Y-%m-%d")),
    ]
    meta_table = doc.add_table(rows=len(meta_rows), cols=2)
    meta_table.style = "Table Grid"
    _lock_table_layout(meta_table, Cm(14.7))
    for row, (label, value) in zip(meta_table.rows, meta_rows):
        label_cell, value_cell = row.cells
        label_cell.width = Cm(3.2)
        value_cell.width = Cm(11.5)
        _set_cell_background(label_cell, DOCX_LIGHT_BLUE)
        label_run = label_cell.paragraphs[0].add_run(label)
        label_run.bold = True
        value_cell.paragraphs[0].add_run(str(value))

    doc.add_paragraph()

    note_p = doc.add_paragraph()
    note_run = note_p.add_run(
        "📌 AI가 생성한 초안입니다. 빨간색으로 표시된 [확인 필요] 부분은 제출 전 실제 정보로 반드시 채워 넣으세요."
    )
    note_run.italic = True
    note_run.font.size = Pt(9.5)
    note_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # 섹션 본문
    comments: list[str] = []
    for idx, (name, text) in enumerate(sections.items(), start=1):
        # 실제 "제목 2" 스타일을 적용해두면 Word의 탐색 창(개요)에 섹션이 잡히고,
        # 나중에 목차(TOC)를 넣어도 자동으로 인식된다 - 색상/테두리는 이후 직접 다시 덮어써서
        # 스타일 적용 여부와 무관하게 지금까지의 디자인을 그대로 유지한다.
        heading_p = doc.add_paragraph(style=_style_or_none(doc, "Heading 2"))
        heading_p.paragraph_format.space_before = Pt(18)
        heading_p.paragraph_format.space_after = Pt(6)
        _add_left_accent_border(heading_p)
        heading_run = heading_p.add_run(f"{idx}. {name}")
        heading_run.bold = True
        heading_run.font.size = Pt(13)
        heading_run.font.color.rgb = DOCX_NAVY

        for para_text in (text or "").split("\n"):
            if not para_text.strip():
                continue
            body_p = doc.add_paragraph()
            body_p.paragraph_format.line_spacing = 1.3
            body_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            _add_body_text(body_p, para_text, comments)

    _add_page_number_footer(doc)
    _attach_comments(doc, comments)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def fetch_docx_bytes(url: str) -> bytes:
    resp = requests.get(url, timeout=20)
    resp.raise_for_status()
    return resp.content


def _insert_paragraph_after(paragraph: Paragraph) -> Paragraph:
    """python-docx엔 '이 문단 뒤에 새 문단 삽입' API가 없어서(항상 문서 끝에만 추가 가능),
    같은 위치에 원소를 만들어 addnext로 옆에 꽂아주는 방식으로 직접 구현한다."""
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def _collect_fill_targets(doc: Document) -> list[dict]:
    """양식에서 내용을 채워 넣을 수 있는 후보 위치를 문단과 '표 셀' 양쪽에서 모두 모은다.

    한국 정부 신청서/사업계획서 양식은 자유서술형 문단보다 [라벨 칸 | 값 칸] 표 구조가
    훨씬 흔한데(예: '지원 사업명' 칸 옆에 값을 적는 칸), 표 셀을 후보에서 빼면 그런 양식은
    채울 수 있는 자리를 아예 못 찾아 매칭 정확도가 크게 떨어진다. 표의 각 행이 2칸 이상이면
    앞 칸을 라벨로 보고 바로 다음 칸을 그 라벨에 대응하는 '값 칸' 후보로 등록한다."""
    targets = []

    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            targets.append({"index": len(targets), "kind": "paragraph", "ref": p, "label": text[:150]})

    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            for c_idx, cell in enumerate(cells):
                label_text = cell.text.strip()
                if not label_text or c_idx + 1 >= len(cells):
                    continue
                value_cell = cells[c_idx + 1]
                if value_cell._tc is cell._tc:
                    continue  # 가로 병합된 칸 - 같은 칸을 라벨로 또 잡지 않는다
                targets.append({
                    "index": len(targets),
                    "kind": "table_cell",
                    "ref": value_cell,
                    "label": f"[표] {label_text[:150]}",
                })

    return targets


def _append_lines_to_cell(cell, text: str, comments: list[str]):
    lines = [line for line in (text or "").split("\n") if line.strip()]
    if not lines:
        return
    # 빈 값 칸(기본 빈 문단 하나만 있는 상태)이면 그 문단을 그대로 재사용해서, 안 그러면
    # 생기는 불필요한 첫 줄 공백(빈 문단 + 새 문단)을 피한다.
    was_empty = not cell.text.strip()
    first_line, rest = lines[0], lines[1:]
    target_p = cell.paragraphs[0] if was_empty else cell.add_paragraph()
    _add_body_text(target_p, first_line, comments)
    for line in rest:
        _add_body_text(cell.add_paragraph(), line, comments)


def fill_docx_template(template_bytes: bytes, sections: dict) -> dict:
    """공고에 첨부된 실제 .docx 신청서 양식을 받아, 그 문서 안의 항목(제목/라벨/표의 값 칸)
    자리에 AI가 작성한 해당 섹션 내용을 직접 삽입한다. build_docx()처럼 새 문서를 만드는
    대신, 원본 양식의 서식(표, 안내문 등)을 그대로 보존한 채 내용만 끼워 넣는 것이 목적이다.

    양식에서 위치를 찾지 못한 항목은 버리지 않고 문서 끝에 별도로 덧붙인다.
    반환값: {"buffer": io.BytesIO, "unmatched": [매칭 안 된 항목명, ...]}
    """
    doc = Document(io.BytesIO(template_bytes))
    targets = _collect_fill_targets(doc)

    mapping = {}
    if targets and sections:
        prompt = TEMPLATE_MAP_PROMPT.format(
            targets_json=json.dumps(
                [{"index": t["index"], "kind": t["kind"], "label": t["label"]} for t in targets],
                ensure_ascii=False,
            ),
            section_names_json=json.dumps(list(sections.keys()), ensure_ascii=False),
        )
        # 표 구조 파악 등 문서 구조 추론이 필요해 단순 추출용 lite 모델 대신 상위 모델을 쓴다.
        response = ai_client.models.generate_content(
            model=WRITING_MODEL,
            contents=prompt,
            config={"response_mime_type": "application/json"},
        )
        mapping = _parse_json_response(response)

    comments: list[str] = []
    unmatched = []
    for name, text in sections.items():
        try:
            idx = int(mapping.get(name, -1))
        except (TypeError, ValueError):
            idx = -1

        if idx < 0 or idx >= len(targets):
            unmatched.append(name)
            continue

        target = targets[idx]
        if target["kind"] == "table_cell":
            _append_lines_to_cell(target["ref"], text, comments)
        else:
            insert_after = target["ref"]
            for line in (text or "").split("\n"):
                if not line.strip():
                    continue
                new_p = _insert_paragraph_after(insert_after)
                _add_body_text(new_p, line, comments)
                insert_after = new_p

    if unmatched:
        doc.add_paragraph()
        note_p = doc.add_paragraph()
        note_run = note_p.add_run(
            "※ 아래 항목은 양식에서 알맞은 위치를 찾지 못해 문서 끝에 추가했습니다. 적절한 위치로 직접 옮겨 주세요."
        )
        note_run.italic = True
        note_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        for name in unmatched:
            heading_p = doc.add_paragraph(style=_style_or_none(doc, "Heading 2"))
            _add_left_accent_border(heading_p)
            heading_run = heading_p.add_run(name)
            heading_run.bold = True
            heading_run.font.color.rgb = DOCX_NAVY
            for line in (sections.get(name) or "").split("\n"):
                if line.strip():
                    body_p = doc.add_paragraph()
                    _add_body_text(body_p, line, comments)

    _attach_comments(doc, comments)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return {"buffer": buf, "unmatched": unmatched}
